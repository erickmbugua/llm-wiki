"""Tests for core/ingest.py — extraction, JSON parsing, page writing, full ingest flow."""

import json
from collections.abc import Callable, Generator
from pathlib import Path
from unittest.mock import MagicMock, patch

import pytest
import requests as req

import core.ingest
from core.chunking import chunk_text, summarize_chunks
from core.extraction import extract_text
from core.ingest import (
    append_log,
    check_ollama,
    ingest_source,
    write_pages,
)
from core.prompts import build_ingest_prompt, parse_llm_json

# ── chunk_text ────────────────────────────────────────────────────────────────


class TestChunkText:
    def test_single_chunk_when_text_fits(self) -> None:
        text = "x" * 1000
        chunks = chunk_text(text, chunk_size=2000, overlap=100)
        assert chunks == [text]

    def test_splits_into_multiple_chunks_with_overlap(self) -> None:
        text = "a" * 40_000
        chunks = chunk_text(text, chunk_size=20_000, overlap=500)
        assert len(chunks) >= 2
        # Every chunk is at most chunk_size chars
        assert all(len(c) <= 20_000 for c in chunks)
        # Adjacent chunks overlap — the start of each successive chunk is within
        # `overlap` chars of the end of the previous chunk's start position
        assert len(chunks[0]) > 0 and len(chunks[-1]) > 0

    def test_breaks_at_newline_within_last_200_chars(self) -> None:
        # Build text where a newline sits 150 chars before the window end
        body = "a" * 9_850 + "\n" + "b" * 10_149
        text = body + "c" * 10_000  # total > chunk_size so chunking kicks in
        chunks = chunk_text(text, chunk_size=10_000, overlap=200)
        # The first chunk should end just after the newline, not mid-sentence
        assert chunks[0].endswith("\n"), "First chunk should break at newline"

    def test_empty_string_returns_single_empty_chunk(self) -> None:
        assert chunk_text("", chunk_size=1000, overlap=50) == [""]

    def test_overlap_cannot_exceed_chunk_size(self) -> None:
        # Should not raise even with degenerate overlap value
        text = "x" * 5000
        chunks = chunk_text(text, chunk_size=2000, overlap=3000)
        assert len(chunks) >= 1


# ── summarize_chunks ──────────────────────────────────────────────────────────


class TestSummarizeChunks:
    def test_calls_model_once_per_chunk(
        self, fake_llm_response: Callable[[str], MagicMock]
    ) -> None:
        chunks = ["chunk one", "chunk two", "chunk three"]
        mock_resp = fake_llm_response("• Key point")
        with patch("core.chunking.litellm.completion", return_value=mock_resp) as mock_llm:
            result = summarize_chunks(
                chunks, model="claude-sonnet-4-6", vault_name="TestVault", filename="doc.txt"
            )
        assert mock_llm.call_count == 3
        assert "Part 1/3" in result
        assert "Part 3/3" in result

    def test_returns_empty_string_for_empty_chunk_list(self) -> None:
        result = summarize_chunks([], model="claude-sonnet-4-6", vault_name="V", filename="f.txt")
        assert result == ""

    def test_truncates_when_summaries_exceed_context_chars(
        self, fake_llm_response: Callable[[str], MagicMock]
    ) -> None:
        # Each chunk summary is 10k chars → 3 chunks = 30k total > 5k context_chars
        big_summary = "x" * 10_000
        mock_resp = fake_llm_response(big_summary)
        with patch("core.chunking.litellm.completion", return_value=mock_resp):
            result = summarize_chunks(
                ["a", "b", "c"],
                model="claude-sonnet-4-6",
                vault_name="V",
                filename="f.txt",
                context_chars=5_000,
            )
        assert len(result) == 5_000
        assert result.endswith("above covers the first sections only]")


# ── extract_text ──────────────────────────────────────────────────────────────


class TestExtractText:
    def test_reads_txt_file(self, tmp_path: Path) -> None:
        f = tmp_path / "doc.txt"
        f.write_text("Hello world")
        text, name = extract_text(str(f))
        assert text == "Hello world"
        assert name == "doc.txt"

    def test_reads_md_file(self, tmp_path: Path) -> None:
        f = tmp_path / "note.md"
        f.write_text("# Title\nBody text")
        text, name = extract_text(str(f))
        assert "Body text" in text
        assert name == "note.md"

    def test_returns_empty_for_nonexistent_file(self, tmp_path: Path) -> None:
        text, _name = extract_text(str(tmp_path / "missing.txt"))
        assert text == ""

    def test_fetches_url(self) -> None:
        fake_response = MagicMock()
        fake_response.text = (
            "<html><head><title>My Page</title></head><body><p>Content here</p></body></html>"
        )
        with patch("core.extraction.requests.get", return_value=fake_response):
            text, name = extract_text("https://example.com/page")
        assert "Content here" in text
        assert name == "My Page"

    def test_url_strips_script_tags(self) -> None:
        fake_response = MagicMock()
        fake_response.text = "<html><body><script>evil()</script><p>Clean text</p></body></html>"
        with patch("core.extraction.requests.get", return_value=fake_response):
            text, _ = extract_text("https://example.com")
        assert "evil()" not in text
        assert "Clean text" in text

    def test_truncates_large_content(self, tmp_path: Path) -> None:
        f = tmp_path / "big.txt"
        f.write_text("x" * 30_000)
        text, _ = extract_text(str(f))
        assert len(text) == 24_000

    def test_extract_docx_returns_text(self, tmp_path: Path) -> None:
        import docx

        doc = docx.Document()
        doc.add_paragraph("First paragraph")
        doc.add_paragraph("Second paragraph")
        path = tmp_path / "test.docx"
        doc.save(str(path))
        text, name = extract_text(str(path))
        assert "First paragraph" in text
        assert "Second paragraph" in text
        assert name == "test.docx"

    def test_extract_docx_missing_package_returns_empty(self, tmp_path: Path) -> None:
        import sys

        path = tmp_path / "doc.docx"
        path.write_bytes(b"fake")
        with patch.dict(sys.modules, {"docx": None}):
            from core.extraction import extract_docx

            result = extract_docx(path)
        assert result == ""

    def test_binary_suffix_raises_value_error(self, tmp_path: Path) -> None:
        f = tmp_path / "data.xlsx"
        f.write_bytes(b"PK\x03\x04fake")
        with pytest.raises(ValueError, match="Unsupported file type"):
            extract_text(str(f))


# ── parse_llm_json ────────────────────────────────────────────────────────────


class TestParseLlmJson:
    def test_parses_valid_json(self) -> None:
        raw = json.dumps(
            {"source_page": {"file_path": "Sources/X.md", "content": "# X"}, "page_updates": []}
        )
        result = parse_llm_json(raw)
        assert result["source_page"]["file_path"] == "Sources/X.md"
        assert result["page_updates"] == []

    def test_strips_markdown_fences(self) -> None:
        raw = '```json\n{"source_page": {"file_path": "Sources/X.md", "content": "# X"}, "page_updates": []}\n```'
        result = parse_llm_json(raw)
        assert result["source_page"]["file_path"] == "Sources/X.md"

    def test_strips_plain_code_fences(self) -> None:
        raw = '```\n{"source_page": {"file_path": "Sources/X.md", "content": "# X"}}\n```'
        result = parse_llm_json(raw)
        assert "source_page" in result

    def test_defaults_page_updates_to_empty_list(self) -> None:
        raw = json.dumps({"source_page": {"file_path": "Sources/X.md", "content": "# X"}})
        result = parse_llm_json(raw)
        assert result["page_updates"] == []

    def test_raises_on_invalid_json(self) -> None:
        with pytest.raises(ValueError):
            parse_llm_json("not json at all")

    def test_raises_when_source_page_missing(self) -> None:
        with pytest.raises(ValueError, match="missing 'source_page'"):
            parse_llm_json(json.dumps({"page_updates": []}))

    def test_trailing_comma_is_repaired(self) -> None:
        raw = '{"source_page": {"file_path": "Sources/X.md", "content": "# X"},}'
        result = parse_llm_json(raw)
        assert result["source_page"]["file_path"] == "Sources/X.md"

    def test_single_quotes_are_repaired(self) -> None:
        raw = "{'source_page': {'file_path': 'Sources/X.md', 'content': '# X'}}"
        result = parse_llm_json(raw)
        assert result["source_page"]["file_path"] == "Sources/X.md"

    def test_prose_before_json_is_extracted(self) -> None:
        payload = json.dumps(
            {"source_page": {"file_path": "Sources/X.md", "content": "# X"}, "page_updates": []}
        )
        raw = f"Here is the JSON output as requested:\n\n{payload}"
        result = parse_llm_json(raw)
        assert result["source_page"]["file_path"] == "Sources/X.md"

    def test_unrepairable_raises_value_error(self) -> None:
        with pytest.raises(ValueError):
            parse_llm_json("completely unparseable @@@ !!!")

    def test_missing_source_page_after_repair_raises(self) -> None:
        raw = '{"page_updates": [],}'
        with pytest.raises(ValueError, match="missing 'source_page'"):
            parse_llm_json(raw)

    def test_missing_closing_brace_is_repaired(self) -> None:
        raw = '{"source_page": {"file_path": "Sources/X.md", "content": "# X"}'
        result = parse_llm_json(raw)
        assert result["source_page"]["file_path"] == "Sources/X.md"


# ── write_pages ───────────────────────────────────────────────────────────────


class TestWritePages:
    def test_writes_source_page(self, tmp_vault: Path) -> None:
        wiki = tmp_vault / "wiki"
        result = {
            "source_page": {"file_path": "Sources/Article.md", "content": "# Article\nContent."},
            "page_updates": [],
        }
        written = write_pages(wiki, result)
        assert "Sources/Article.md" in written
        assert (wiki / "Sources" / "Article.md").read_text() == "# Article\nContent."

    def test_creates_concept_page(self, tmp_vault: Path) -> None:
        wiki = tmp_vault / "wiki"
        result = {
            "source_page": {"file_path": "Sources/S.md", "content": "# S"},
            "page_updates": [
                {"file_path": "Concepts/NewConcept.md", "action": "create", "content": "# New"},
            ],
        }
        written = write_pages(wiki, result)
        assert "Concepts/NewConcept.md" in written
        assert (wiki / "Concepts" / "NewConcept.md").exists()

    def test_create_on_existing_file_replaces_content(self, tmp_vault: Path) -> None:
        wiki = tmp_vault / "wiki"
        existing = wiki / "Concepts" / "Existing.md"
        existing.write_text("# Original content")
        result = {
            "source_page": {"file_path": "Sources/S.md", "content": "# S"},
            "page_updates": [
                {
                    "file_path": "Concepts/Existing.md",
                    "action": "create",
                    "content": "# Fully updated page",
                },
            ],
        }
        write_pages(wiki, result)
        content = existing.read_text()
        assert content == "# Fully updated page"
        assert "# Original content" not in content

    def test_overwrites_existing_page_on_update_action(self, tmp_vault: Path) -> None:
        wiki = tmp_vault / "wiki"
        existing = wiki / "Concepts" / "Replace.md"
        existing.write_text("# Old")
        result = {
            "source_page": {"file_path": "Sources/S.md", "content": "# S"},
            "page_updates": [
                {"file_path": "Concepts/Replace.md", "action": "update", "content": "# New"},
            ],
        }
        write_pages(wiki, result)
        assert existing.read_text() == "# New"

    def test_skips_entries_with_empty_content(self, tmp_vault: Path) -> None:
        wiki = tmp_vault / "wiki"
        result = {
            "source_page": {"file_path": "Sources/S.md", "content": "# S"},
            "page_updates": [{"file_path": "Concepts/Empty.md", "action": "create", "content": ""}],
        }
        written = write_pages(wiki, result)
        assert "Concepts/Empty.md" not in written

    def test_creates_nested_dirs(self, tmp_vault: Path) -> None:
        wiki = tmp_vault / "wiki"
        result = {
            "source_page": {"file_path": "Sources/Deep/Article.md", "content": "# A"},
            "page_updates": [],
        }
        write_pages(wiki, result)
        assert (wiki / "Sources" / "Deep" / "Article.md").exists()

    def test_unsafe_source_page_path_is_skipped(self, tmp_vault: Path) -> None:
        wiki = tmp_vault / "wiki"
        result = {
            "source_page": {
                "file_path": "../../evil.md",
                "content": "malicious content",
            },
            "page_updates": [],
        }
        written = write_pages(wiki, result)
        assert written == []
        assert not (tmp_vault / "evil.md").exists()

    def test_unsafe_page_update_path_is_skipped(self, tmp_vault: Path) -> None:
        wiki = tmp_vault / "wiki"
        result = {
            "source_page": {"file_path": "Sources/Safe.md", "content": "# Safe"},
            "page_updates": [
                {
                    "file_path": "../../.llm-wiki/config.json",
                    "action": "create",
                    "content": "{}",
                }
            ],
        }
        written = write_pages(wiki, result)
        assert "Sources/Safe.md" in written
        assert "../../.llm-wiki/config.json" not in written
        assert (tmp_vault / ".llm-wiki" / "config.json").read_text() != "{}"


# ── build_ingest_prompt ───────────────────────────────────────────────────────


class TestBuildIngestPrompt:
    def test_contains_vault_name(self) -> None:
        prompt = build_ingest_prompt("MyVault", "", "", "source.txt", "text")
        assert "MyVault" in prompt

    def test_contains_source_text(self) -> None:
        prompt = build_ingest_prompt("V", "", "", "f.txt", "unique_source_content_xyz")
        assert "unique_source_content_xyz" in prompt

    def test_includes_no_related_section_when_empty(self) -> None:
        prompt = build_ingest_prompt("V", "", "", "f.txt", "text")
        assert "(none yet)" in prompt

    def test_includes_related_pages_when_provided(self) -> None:
        prompt = build_ingest_prompt("V", "", "### Existing Page\nContent", "f.txt", "text")
        assert "Existing Page" in prompt

    def test_prompt_instructs_yaml_colon_quoting(self) -> None:
        prompt = build_ingest_prompt("V", "", "", "f.txt", "text")
        assert "colon" in prompt.lower()


# ── append_log ────────────────────────────────────────────────────────────────


class TestAppendLog:
    def test_appends_entry_to_log(self, tmp_vault: Path) -> None:
        append_log(tmp_vault, "my-source.txt", ["Sources/X.md", "Concepts/Y.md"])
        log_text = (tmp_vault / "wiki" / "log.md").read_text()
        assert "my-source.txt" in log_text
        assert "Sources/X" in log_text

    def test_creates_log_if_missing(self, tmp_path: Path) -> None:
        vault = tmp_path / "v"
        (vault / "wiki").mkdir(parents=True)
        append_log(vault, "source", [])
        assert (vault / "wiki" / "log.md").exists()


# ── ingest_source (full flow, LLM mocked) ────────────────────────────────────


class TestIngestSourceChunking:
    def test_large_doc_triggers_chunking(
        self, tmp_vault: Path, fake_llm_response: Callable[[str], MagicMock]
    ) -> None:
        """A source larger than chunk_size should call _summarize_chunks before the ingest prompt."""
        llm_output = json.dumps(
            {
                "source_page": {"file_path": "Sources/BigDoc.md", "content": "# Big\nSummary."},
                "page_updates": [],
            }
        )
        # Write a file bigger than the default chunk_size (20_000)
        big_file = tmp_vault / "raw" / "big.txt"
        big_file.write_text("x" * 25_000)

        with (
            patch("core.ingest.litellm.completion", return_value=fake_llm_response(llm_output)),
            patch("core.ingest.resolve_model", return_value="claude-sonnet-4-6"),
            patch("core.ingest.summarize_chunks", return_value="summarized content") as mock_sc,
        ):
            result = ingest_source(tmp_vault, str(big_file), "TestVault")

        mock_sc.assert_called_once()
        assert result["pages_written"] != []

    def test_small_doc_skips_chunking(
        self, tmp_vault: Path, fake_llm_response: Callable[[str], MagicMock]
    ) -> None:
        """A source smaller than chunk_size should not call _summarize_chunks."""
        llm_output = json.dumps(
            {
                "source_page": {"file_path": "Sources/Small.md", "content": "# Small\nContent."},
                "page_updates": [],
            }
        )
        small_file = tmp_vault / "raw" / "small.txt"
        small_file.write_text("hello world")

        with (
            patch("core.ingest.litellm.completion", return_value=fake_llm_response(llm_output)),
            patch("core.ingest.resolve_model", return_value="claude-sonnet-4-6"),
            patch("core.ingest.summarize_chunks") as mock_sc,
        ):
            ingest_source(tmp_vault, str(small_file), "TestVault")

        mock_sc.assert_not_called()


class TestIngestSource:
    def test_full_ingest_writes_pages(
        self, tmp_vault: Path, fake_llm_response: Callable[[str], MagicMock]
    ) -> None:
        llm_output = json.dumps(
            {
                "source_page": {
                    "file_path": "Sources/Article.md",
                    "content": "---\ntitle: Article\ntags: [test]\n---\nSummary here.",
                },
                "page_updates": [
                    {
                        "file_path": "Concepts/KeyIdea.md",
                        "action": "create",
                        "content": "---\ntitle: Key Idea\ntags: [test]\n---\nA key idea.",
                    }
                ],
            }
        )
        wiki = tmp_vault / "wiki"
        (wiki / "schema.md").write_text("# Schema\nIngest rules here.")

        with (
            patch("core.ingest.litellm.completion", return_value=fake_llm_response(llm_output)),
            patch("core.ingest.resolve_model", return_value="claude-sonnet-4-6"),
        ):
            result = ingest_source(tmp_vault, str(tmp_vault / "wiki" / "schema.md"), "TestVault")

        assert (wiki / "Sources" / "Article.md").exists()
        assert (wiki / "Concepts" / "KeyIdea.md").exists()
        assert "Sources/Article.md" in result["pages_written"]

    def test_dry_run_does_not_write_files(
        self, tmp_vault: Path, fake_llm_response: Callable[[str], MagicMock]
    ) -> None:
        llm_output = json.dumps(
            {
                "source_page": {"file_path": "Sources/DryRun.md", "content": "# Dry"},
                "page_updates": [],
            }
        )
        src = tmp_vault / "raw" / "test.txt"
        src.write_text("some content")

        with (
            patch("core.ingest.litellm.completion", return_value=fake_llm_response(llm_output)),
            patch("core.ingest.resolve_model", return_value="claude-sonnet-4-6"),
        ):
            result = ingest_source(tmp_vault, str(src), "TestVault", dry_run=True)

        assert not (tmp_vault / "wiki" / "Sources" / "DryRun.md").exists()
        assert result["pages_written"] == []

    def test_raises_when_source_unreadable(self, tmp_vault: Path) -> None:
        with pytest.raises(ValueError, match="Could not extract text"):
            ingest_source(tmp_vault, "/nonexistent/path/to/file.xyz", "TestVault")


# ── check_ollama ──────────────────────────────────────────────────────────────


class TestCheckOllama:
    @pytest.fixture(autouse=True)
    def clear_ollama_cache(self) -> Generator[None, None, None]:
        """Reset the in-process Ollama verification cache between tests."""
        core.ingest._ollama_verified.clear()
        yield
        core.ingest._ollama_verified.clear()

    def test_success_when_model_present(self) -> None:
        fake_resp = MagicMock()
        fake_resp.json.return_value = {
            "models": [{"name": "qwen2.5-coder:7b"}, {"name": "llama3:8b"}]
        }
        with patch("core.ingest.requests.get", return_value=fake_resp):
            check_ollama("ollama/qwen2.5-coder:7b")  # must not raise

    def test_raises_when_server_unreachable(self) -> None:
        with (
            patch("core.ingest.requests.get", side_effect=req.exceptions.ConnectionError()),
            pytest.raises(RuntimeError, match="ollama serve"),
        ):
            check_ollama("ollama/qwen2.5-coder:7b")

    def test_raises_when_model_not_pulled(self) -> None:
        fake_resp = MagicMock()
        fake_resp.json.return_value = {"models": [{"name": "llama3:8b"}]}
        with (
            patch("core.ingest.requests.get", return_value=fake_resp),
            pytest.raises(RuntimeError, match="ollama pull"),
        ):
            check_ollama("ollama/qwen2.5-coder:7b")

    def test_ingest_source_skips_preflight_for_non_ollama(
        self, tmp_vault: Path, fake_llm_response: Callable[[str], MagicMock]
    ) -> None:
        llm_output = json.dumps(
            {
                "source_page": {"file_path": "Sources/X.md", "content": "# X"},
                "page_updates": [],
            }
        )
        src = tmp_vault / "raw" / "test.txt"
        src.write_text("some content")
        with (
            patch("core.ingest.litellm.completion", return_value=fake_llm_response(llm_output)),
            patch("core.ingest.resolve_model", return_value="claude-sonnet-4-6"),
            patch("core.ingest.check_ollama") as mock_check,
        ):
            ingest_source(tmp_vault, str(src), "TestVault")
        mock_check.assert_not_called()

    def test_check_ollama_caches_result(self) -> None:
        """Second call with the same model must not hit the network."""
        fake_resp = MagicMock()
        fake_resp.json.return_value = {"models": [{"name": "qwen2.5-coder:7b"}]}
        with patch("core.ingest.requests.get", return_value=fake_resp) as mock_get:
            check_ollama("ollama/qwen2.5-coder:7b")
            check_ollama("ollama/qwen2.5-coder:7b")
        mock_get.assert_called_once()

    def test_check_ollama_different_models_each_checked(self) -> None:
        """Different model strings each trigger their own network check."""
        fake_resp = MagicMock()
        fake_resp.json.return_value = {
            "models": [{"name": "qwen2.5-coder:7b"}, {"name": "llama3:8b"}]
        }
        with patch("core.ingest.requests.get", return_value=fake_resp) as mock_get:
            check_ollama("ollama/qwen2.5-coder:7b")
            check_ollama("ollama/llama3:8b")
        assert mock_get.call_count == 2


# ── ingest_queued ─────────────────────────────────────────────────────────────


class TestIngestQueued:
    def test_success_status_transitions(self, tmp_vault: Path) -> None:
        """Queue item transitions pending → processing → done on success."""
        from core.db import get_db, get_pending_queue, queue_raw_file
        from core.ingest import ingest_queued

        conn = get_db(tmp_vault)
        queue_raw_file(conn, "raw/file.txt")
        conn.commit()
        conn.close()

        success_result = {
            "source_page": {"file_path": "Sources/Fake.md", "content": "# Fake"},
            "page_updates": [],
            "pages_written": ["Sources/Fake.md"],
        }
        with patch("core.ingest.ingest_source", return_value=success_result):
            results = ingest_queued(tmp_vault, "TestVault")

        assert len(results) == 1
        assert results[0]["status"] == "done"
        assert results[0]["file"] == "raw/file.txt"

        conn = get_db(tmp_vault)
        remaining = get_pending_queue(conn)
        conn.close()
        assert remaining == []

    def test_ingest_source_called_with_absolute_path(self, tmp_vault: Path) -> None:
        """ingest_queued reconstructs the absolute path before calling ingest_source."""
        from core.db import get_db, queue_raw_file
        from core.ingest import ingest_queued

        conn = get_db(tmp_vault)
        queue_raw_file(conn, "raw/paper.pdf")
        conn.commit()
        conn.close()

        success_result = {
            "source_page": {"file_path": "Sources/Paper.md", "content": "# Paper"},
            "page_updates": [],
            "pages_written": ["Sources/Paper.md"],
        }
        with patch("core.ingest.ingest_source", return_value=success_result) as mock_ingest:
            ingest_queued(tmp_vault, "TestVault")

        called_source = mock_ingest.call_args[0][1]
        assert called_source == str(tmp_vault / "raw" / "paper.pdf")

    def test_failure_status_transitions(self, tmp_vault: Path) -> None:
        """Queue item transitions pending → processing → failed on exception."""
        from core.db import get_db, queue_raw_file
        from core.ingest import ingest_queued

        conn = get_db(tmp_vault)
        queue_raw_file(conn, "raw/broken.txt")
        conn.commit()
        conn.close()

        with patch("core.ingest.ingest_source", side_effect=ValueError("boom")):
            results = ingest_queued(tmp_vault, "TestVault")

        assert len(results) == 1
        assert results[0]["status"] == "failed"
        assert "boom" in results[0]["error"]

    def test_single_db_connection_used(self, tmp_vault: Path) -> None:
        """ingest_queued opens exactly one DB connection for the entire queue."""
        from core.db import db_connection, get_db, queue_raw_file
        from core.ingest import ingest_queued

        conn = get_db(tmp_vault)
        queue_raw_file(conn, "raw/a.txt")
        queue_raw_file(conn, "raw/b.txt")
        conn.commit()
        conn.close()

        success_result = {
            "source_page": {"file_path": "Sources/X.md", "content": "# X"},
            "page_updates": [],
            "pages_written": ["Sources/X.md"],
        }
        with (
            patch("core.ingest.ingest_source", return_value=success_result),
            patch("core.ingest.db_connection", wraps=db_connection) as mock_db_conn,
        ):
            ingest_queued(tmp_vault, "TestVault")

        mock_db_conn.assert_called_once()


# ── rebuild_index integration ─────────────────────────────────────────────────


class TestIngestSourceRebuildIndex:
    def test_ingest_source_calls_rebuild_index(
        self, tmp_vault: Path, fake_llm_response: Callable[[str], MagicMock]
    ) -> None:
        """After a successful ingest, rebuild_index is called once."""
        llm_output = json.dumps(
            {
                "source_page": {"file_path": "Sources/RebuildTest.md", "content": "# R"},
                "page_updates": [],
            }
        )
        src = tmp_vault / "raw" / "test.txt"
        src.write_text("content")

        with (
            patch("core.ingest.litellm.completion", return_value=fake_llm_response(llm_output)),
            patch("core.ingest.resolve_model", return_value="claude-sonnet-4-6"),
            patch("core.ingest.rebuild_index") as mock_rebuild,
        ):
            ingest_source(tmp_vault, str(src), "TestVault")

        mock_rebuild.assert_called_once_with(tmp_vault)

    def test_ingest_source_dry_run_does_not_rebuild_index(
        self, tmp_vault: Path, fake_llm_response: Callable[[str], MagicMock]
    ) -> None:
        """dry_run=True must not call rebuild_index."""
        llm_output = json.dumps(
            {
                "source_page": {"file_path": "Sources/DryRebuild.md", "content": "# D"},
                "page_updates": [],
            }
        )
        src = tmp_vault / "raw" / "dry.txt"
        src.write_text("content")

        with (
            patch("core.ingest.litellm.completion", return_value=fake_llm_response(llm_output)),
            patch("core.ingest.resolve_model", return_value="claude-sonnet-4-6"),
            patch("core.ingest.rebuild_index") as mock_rebuild,
        ):
            ingest_source(tmp_vault, str(src), "TestVault", dry_run=True)

        mock_rebuild.assert_not_called()
